import pandas as pd
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

# Hàm đọc dữ liệu App CSKH
def load_app_data(file):
    df = pd.read_excel(file, skiprows=2)
    df = df.rename(columns=lambda x: x.strip())
    df = df.dropna(subset=['Điện lực'])
    df_app = df[df['Điện lực'] != 'Tổng cộng'].copy()
    df_app_total = df[df['Điện lực'] == 'Tổng cộng'].copy()
    for col in ['Số lượng KH quản lý', 'Số lượng đã thực hiện App', 'Tỷ lệ thực hiện qua App']:
        df_app[col] = pd.to_numeric(df_app[col], errors='coerce')
        df_app_total[col] = pd.to_numeric(df_app_total[col], errors='coerce')
    return df_app, df_app_total

# Hàm đọc dữ liệu Giải quyết đúng hạn
def load_time_data(file):
    df = pd.read_excel(file, skiprows=3)
    df = df.rename(columns=lambda x: x.strip())
    df = df.dropna(subset=['Điện lực'])
    df_time = df[df['Điện lực'] != 'Tổng cộng'].copy()
    df_time_total = df[df['Điện lực'] == 'Tổng cộng'].copy()
    for col in ['Số yêu cầu chuyển xử lý', 'Số lượng phiếu giải quyết trễ hạn', 'Tỷ lệ trễ hạn']:
        df_time[col] = pd.to_numeric(df_time[col], errors='coerce')
        df_time_total[col] = pd.to_numeric(df_time_total[col], errors='coerce')
    return df_time, df_time_total

# Hàm vẽ biểu đồ cột với giá trị hiển thị 6 số thập phân
def plot_bar(df, x_col, y_col, title, y_label, percent=False, color='blue'):
    fig, ax = plt.subplots(figsize=(8,5))
    values = df[y_col]*100 if percent else df[y_col]
    bars = ax.bar(df[x_col], values, color=color)
    ax.set_title(title)
    ax.set_ylabel(y_label)
    ax.set_xticklabels(df[x_col], rotation=45, ha='right')
    for bar, val in zip(bars, values):
        height = bar.get_height()
        ax.annotate(f'{val:.6f}%',
                    xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3),
                    textcoords="offset points",
                    ha='center', va='bottom', fontsize=8)
    plt.tight_layout()
    return fig

# Hàm lưu ảnh matplotlib vào BytesIO
def fig_to_bytes(fig):
    img_bytes = BytesIO()
    fig.savefig(img_bytes, format='png')
    img_bytes.seek(0)
    plt.close(fig)
    return img_bytes

# Hàm xuất Word tab App CSKH kèm biểu đồ
def export_word_app(app_df, app_total, app_info, top3, bottom3, fig_total, fig_top3, fig_bot3):
    doc = Document()
    doc.add_heading('BÁO CÁO APP CSKH', 0)

    # Đánh giá tổng quan
    doc.add_heading('Đánh giá tổng quan kết quả tổng cộng', level=1)
    if not app_total.empty:
        row = app_total.iloc[0]
        doc.add_paragraph(f"Tổng số khách hàng quản lý: {int(row['Số lượng KH quản lý']):,}")
        doc.add_paragraph(f"Tổng số khách hàng đã thực hiện App: {int(row['Số lượng đã thực hiện App']):,}")
        doc.add_paragraph(f"Tỷ lệ thực hiện qua App: {row['Tỷ lệ thực hiện qua App']*100:.6f}%")
    else:
        doc.add_paragraph(f"Tổng số khách hàng quản lý: {app_info['total_kh']:,}")
        doc.add_paragraph(f"Tổng số khách hàng đã thực hiện App: {app_info['total_app']:,}")
        doc.add_paragraph(f"Tỷ lệ thực hiện qua App trung bình: {app_info['avg_rate']*100:.6f}%")

    # Biểu đồ tổng hợp
    doc.add_heading('Biểu đồ kết quả thực hiện tổng hợp', level=1)
    img_total = fig_to_bytes(fig_total)
    doc.add_picture(img_total, width=Inches(6))

    # Bảng dữ liệu tổng hợp
    doc.add_heading('Bảng dữ liệu tổng hợp', level=1)
    columns = ['STT', 'Điện lực', 'Số lượng KH quản lý', 'Số lượng đã thực hiện App', 'Tỷ lệ thực hiện qua App (%)']
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    for _, row in app_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(int(row['STT']))
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số lượng KH quản lý']))
        row_cells[3].text = str(int(row['Số lượng đã thực hiện App']))
        row_cells[4].text = f"{row['Tỷ lệ thực hiện qua App']*100:.6f}%"
    if not app_total.empty:
        row = app_total.iloc[0]
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số lượng KH quản lý']))
        row_cells[3].text = str(int(row['Số lượng đã thực hiện App']))
        row_cells[4].text = f"{row['Tỷ lệ thực hiện qua App']*100:.6f}%"

    # Biểu đồ top 3
    doc.add_heading('Top 3 điện lực có tỷ lệ thực hiện cao nhất', level=1)
    img_top3 = fig_to_bytes(fig_top3)
    doc.add_picture(img_top3, width=Inches(6))
    # Bảng top 3
    doc.add_heading('Bảng dữ liệu Top 3', level=2)
    table_top3 = doc.add_table(rows=1, cols=len(columns))
    table_top3.style = 'Light List Accent 1'
    hdr_cells = table_top3.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    for _, row in top3.iterrows():
        row_cells = table_top3.add_row().cells
        row_cells[0].text = str(int(row['STT']))
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số lượng KH quản lý']))
        row_cells[3].text = str(int(row['Số lượng đã thực hiện App']))
        row_cells[4].text = f"{row['Tỷ lệ thực hiện qua App']*100:.6f}%"

    # Biểu đồ bottom 3
    doc.add_heading('Bottom 3 điện lực có tỷ lệ thực hiện thấp nhất', level=1)
    img_bot3 = fig_to_bytes(fig_bot3)
    doc.add_picture(img_bot3, width=Inches(6))
    # Bảng bottom 3
    doc.add_heading('Bảng dữ liệu Bottom 3', level=2)
    table_bot3 = doc.add_table(rows=1, cols=len(columns))
    table_bot3.style = 'Light List Accent 1'
    hdr_cells = table_bot3.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    for _, row in bottom3.iterrows():
        row_cells = table_bot3.add_row().cells
        row_cells[0].text = str(int(row['STT']))
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số lượng KH quản lý']))
        row_cells[3].text = str(int(row['Số lượng đã thực hiện App']))
        row_cells[4].text = f"{row['Tỷ lệ thực hiện qua App']*100:.6f}%"

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# Hàm xuất Word tab Giải quyết đúng hạn kèm biểu đồ
def export_word_time(time_df, time_total, time_info, top3, bottom3, fig_total, fig_top3, fig_bot3):
    doc = Document()
    doc.add_heading('BÁO CÁO GIẢI QUYẾT ĐÚNG THỜI GIAN', 0)

    # Đánh giá tổng quan
    doc.add_heading('Đánh giá tổng quan kết quả tổng cộng', level=1)
    if not time_total.empty:
        row = time_total.iloc[0]
        doc.add_paragraph(f"Tổng số yêu cầu chuyển xử lý: {int(row['Số yêu cầu chuyển xử lý']):,}")
        doc.add_paragraph(f"Tổng số phiếu giải quyết trễ hạn: {int(row['Số lượng phiếu giải quyết trễ hạn']):,}")
        doc.add_paragraph(f"Tỷ lệ trễ hạn: {row['Tỷ lệ trễ hạn']*100:.6f}%")
    else:
        doc.add_paragraph(f"Tổng số yêu cầu chuyển xử lý: {time_info['total_req']:,}")
        doc.add_paragraph(f"Tổng số phiếu giải quyết trễ hạn: {time_info['total_late']:,}")
        doc.add_paragraph(f"Tỷ lệ trễ hạn trung bình: {time_info['avg_late']*100:.6f}%")

    # Biểu đồ tổng hợp
    doc.add_heading('Biểu đồ kết quả thực hiện tổng hợp', level=1)
    img_total = fig_to_bytes(fig_total)
    doc.add_picture(img_total, width=Inches(6))

    # Bảng dữ liệu tổng hợp
    doc.add_heading('Bảng dữ liệu tổng hợp', level=1)
    columns = ['STT', 'Điện lực', 'Số yêu cầu chuyển xử lý', 'Số lượng phiếu giải quyết trễ hạn', 'Tỷ lệ trễ hạn (%)']
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    for _, row in time_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(int(row['STT']))
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số yêu cầu chuyển xử lý']))
        row_cells[3].text = str(int(row['Số lượng phiếu giải quyết trễ hạn']))
        row_cells[4].text = f"{row['Tỷ lệ trễ hạn']*100:.6f}%"
    if not time_total.empty:
        row = time_total.iloc[0]
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số yêu cầu chuyển xử lý']))
        row_cells[3].text = str(int(row['Số lượng phiếu giải quyết trễ hạn']))
        row_cells[4].text = f"{row['Tỷ lệ trễ hạn']*100:.6f}%"

    # Biểu đồ top 3
    doc.add_heading('Top 3 điện lực có tỷ lệ trễ hạn cao nhất', level=1)
    img_top3 = fig_to_bytes(fig_top3)
    doc.add_picture(img_top3, width=Inches(6))
    # Bảng top 3
    doc.add_heading('Bảng dữ liệu Top 3', level=2)
    table_top3 = doc.add_table(rows=1, cols=len(columns))
    table_top3.style = 'Light List Accent 1'
    hdr_cells = table_top3.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    for _, row in top3.iterrows():
        row_cells = table_top3.add_row().cells
        row_cells[0].text = str(int(row['STT']))
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số yêu cầu chuyển xử lý']))
        row_cells[3].text = str(int(row['Số lượng phiếu giải quyết trễ hạn']))
        row_cells[4].text = f"{row['Tỷ lệ trễ hạn']*100:.6f}%"

    # Biểu đồ bottom 3
    doc.add_heading('Bottom 3 điện lực có tỷ lệ trễ hạn thấp nhất', level=1)
    img_bot3 = fig_to_bytes(fig_bot3)
    doc.add_picture(img_bot3, width=Inches(6))
    # Bảng bottom 3
    doc.add_heading('Bảng dữ liệu Bottom 3', level=2)
    table_bot3 = doc.add_table(rows=1, cols=len(columns))
    table_bot3.style = 'Light List Accent 1'
    hdr_cells = table_bot3.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    for _, row in bottom3.iterrows():
        row_cells = table_bot3.add_row().cells
        row_cells[0].text = str(int(row['STT']))
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số yêu cầu chuyển xử lý']))
        row_cells[3].text = str(int(row['Số lượng phiếu giải quyết trễ hạn']))
        row_cells[4].text = f"{row['Tỷ lệ trễ hạn']*100:.6f}%"

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# Streamlit app
st.title("Báo cáo App CSKH và Giải quyết yêu cầu")

tab1, tab2 = st.tabs(["App CSKH", "Giải quyết đúng hạn"])

with tab1:
    file_app = st.file_uploader("Chọn file báo cáo App CSKH", type=["xlsx"], key="app")
    if file_app:
        df_app, df_app_total = load_app_data(file_app)
        df_app_show = pd.concat([df_app, df_app_total], ignore_index=True)
        df_app_show['Tỷ lệ thực hiện qua App (%)'] = df_app_show['Tỷ lệ thực hiện qua App'].apply(lambda x: f"{x*100:.6f}")
        st.dataframe(df_app_show.drop(columns=['Tỷ lệ thực hiện qua App']), hide_index=True)

        total_kh = df_app['Số lượng KH quản lý'].sum()
        total_app = df_app['Số lượng đã thực hiện App'].sum()
        avg_rate = df_app['Tỷ lệ thực hiện qua App'].mean()

        # Đánh giá tổng quan kết quả tổng cộng
        if not df_app_total.empty:
            total_row = df_app_total.iloc[0]
            st.markdown("### Đánh giá tổng quan kết quả tổng cộng")
            st.write(f"- Tổng số khách hàng quản lý: {int(total_row['Số lượng KH quản lý']):,}")
            st.write(f"- Tổng số khách hàng đã thực hiện App: {int(total_row['Số lượng đã thực hiện App']):,}")
            st.write(f"- Tỷ lệ thực hiện qua App: {total_row['Tỷ lệ thực hiện qua App']*100:.6f}%")
        else:
            st.markdown("### Đánh giá tổng quan kết quả tổng cộng")
            st.write(f"- Tổng số khách hàng quản lý: {total_kh:,}")
            st.write(f"- Tổng số khách hàng đã thực hiện App: {total_app:,}")
            st.write(f"- Tỷ lệ thực hiện qua App trung bình: {avg_rate*100:.6f}%")

        # Biểu đồ kết quả tổng hợp toàn bộ điện lực
        fig_total = plot_bar(df_app, 'Điện lực', 'Tỷ lệ thực hiện qua App', "Kết quả thực hiện qua App CSKH", "Tỷ lệ (%)", percent=True, color='mediumseagreen')
        st.pyplot(fig_total)

        # Top 3 và Bottom 3
        top3 = df_app.nlargest(3, 'Tỷ lệ thực hiện qua App')
        bottom3 = df_app.nsmallest(3, 'Tỷ lệ thực hiện qua App')

        st.subheader("Top 3 điện lực có tỷ lệ thực hiện cao nhất")
        st.dataframe(top3[['STT', 'Điện lực', 'Số lượng KH quản lý', 'Số lượng đã thực hiện App', 'Tỷ lệ thực hiện qua App']].assign(**{'Tỷ lệ thực hiện qua App (%)': top3['Tỷ lệ thực hiện qua App'].apply(lambda x: f"{x*100:.6f}")}).drop(columns=['Tỷ lệ thực hiện qua App']))

        st.subheader("Bottom 3 điện lực có tỷ lệ thực hiện thấp nhất")
        st.dataframe(bottom3[['STT', 'Điện lực', 'Số lượng KH quản lý', 'Số lượng đã thực hiện App', 'Tỷ lệ thực hiện qua App']].assign(**{'Tỷ lệ thực hiện qua App (%)': bottom3['Tỷ lệ thực hiện qua App'].apply(lambda x: f"{x*100:.6f}")}).drop(columns=['Tỷ lệ thực hiện qua App']))

        fig_top3 = plot_bar(top3, 'Điện lực', 'Tỷ lệ thực hiện qua App', "Top 3 App CSKH", "Tỷ lệ (%)", percent=True, color='royalblue')
        st.pyplot(fig_top3)

        fig_bot3 = plot_bar(bottom3, 'Điện lực', 'Tỷ lệ thực hiện qua App', "Bottom 3 App CSKH", "Tỷ lệ (%)", percent=True, color='orange')
        st.pyplot(fig_bot3)

        if st.button("Tải báo cáo Word tab App CSKH"):
            word_file = export_word_app(df_app, df_app_total, {'total_kh': total_kh, 'total_app': total_app, 'avg_rate': avg_rate}, top3, bottom3, fig_total, fig_top3, fig_bot3)
            st.download_button("Tải file Word", data=word_file, file_name="Bao_cao_AppCSKH.docx")

with tab2:
    file_time = st.file_uploader("Chọn file báo cáo trễ hạn", type=["xlsx"], key="time")
    if file_time:
        df_time, df_time_total = load_time_data(file_time)
        df_time_show = pd.concat([df_time, df_time_total], ignore_index=True)
        df_time_show['Tỷ lệ trễ hạn (%)'] = df_time_show['Tỷ lệ trễ hạn'].apply(lambda x: f"{x*100:.6f}")
        st.dataframe(df_time_show.drop(columns=['Tỷ lệ trễ hạn']), hide_index=True)

        total_req = df_time['Số yêu cầu chuyển xử lý'].sum()
        total_late = df_time['Số lượng phiếu giải quyết trễ hạn'].sum()
        avg_late = df_time['Tỷ lệ trễ hạn'].mean()

        # Đánh giá tổng quan kết quả tổng cộng
        if not df_time_total.empty:
            total_row = df_time_total.iloc[0]
            st.markdown("### Đánh giá tổng quan kết quả tổng cộng")
            st.write(f"- Tổng số yêu cầu chuyển xử lý: {int(total_row['Số yêu cầu chuyển xử lý']):,}")
            st.write(f"- Tổng số phiếu giải quyết trễ hạn: {int(total_row['Số lượng phiếu giải quyết trễ hạn']):,}")
            st.write(f"- Tỷ lệ trễ hạn: {total_row['Tỷ lệ trễ hạn']*100:.6f}%")
        else:
            st.markdown("### Đánh giá tổng quan kết quả tổng cộng")
            st.write(f"- Tổng số yêu cầu chuyển xử lý: {total_req:,}")
            st.write(f"- Tổng số phiếu giải quyết trễ hạn: {total_late:,}")
            st.write(f"- Tỷ lệ trễ hạn trung bình: {avg_late*100:.6f}%")

        # Biểu đồ kết quả tổng hợp toàn bộ điện lực
        fig_total = plot_bar(df_time, 'Điện lực', 'Tỷ lệ trễ hạn', "Kết quả giải quyết trễ hạn", "Tỷ lệ (%)", percent=True, color='mediumvioletred')
        st.pyplot(fig_total)

        # Top 3 và Bottom 3
        top3 = df_time.nlargest(3, 'Tỷ lệ trễ hạn')
        bottom3 = df_time.nsmallest(3, 'Tỷ lệ trễ hạn')

        st.subheader("Top 3 điện lực có tỷ lệ trễ hạn cao nhất")
        st.dataframe(top3[['STT', 'Điện lực', 'Số yêu cầu chuyển xử lý', 'Số lượng phiếu giải quyết trễ hạn', 'Tỷ lệ trễ hạn']].assign(**{'Tỷ lệ trễ hạn (%)': top3['Tỷ lệ trễ hạn'].apply(lambda x: f"{x*100:.6f}")}).drop(columns=['Tỷ lệ trễ hạn']))

        st.subheader("Bottom 3 điện lực có tỷ lệ trễ hạn thấp nhất")
        st.dataframe(bottom3[['STT', 'Điện lực', 'Số yêu cầu chuyển xử lý', 'Số lượng phiếu giải quyết trễ hạn', 'Tỷ lệ trễ hạn']].assign(**{'Tỷ lệ trễ hạn (%)': bottom3['Tỷ lệ trễ hạn'].apply(lambda x: f"{x*100:.6f}")}).drop(columns=['Tỷ lệ trễ hạn']))

        fig_top3 = plot_bar(top3, 'Điện lực', 'Tỷ lệ trễ hạn', "Top 3 trễ hạn", "Tỷ lệ (%)", percent=True, color='crimson')
        st.pyplot(fig_top3)

        fig_bot3 = plot_bar(bottom3, 'Điện lực', 'Tỷ lệ trễ hạn', "Bottom 3 trễ hạn", "Tỷ lệ (%)", percent=True, color='goldenrod')
        st.pyplot(fig_bot3)

        if st.button("Tải báo cáo Word tab Giải quyết đúng hạn"):
            word_file = export_word_time(df_time, df_time_total, {'total_req': total_req, 'total_late': total_late, 'avg_late': avg_late}, top3, bottom3, fig_total, fig_top3, fig_bot3)
            st.download_button("Tải file Word", data=word_file, file_name="Bao_cao_TreHan.docx")
